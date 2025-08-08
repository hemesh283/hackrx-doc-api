# doc-insight-api/main.py
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv
import fitz  # PyMuPDF
import docx
import email
import email.policy
import os
import tempfile
import tiktoken
import json
import logging
from openai import OpenAI
from sqlalchemy.orm import Session
from db import SessionLocal
from models import Document, QueryLog
import uuid
from typing import Dict, List, Optional
import os
import uvicorn

# Load environment variables
load_dotenv()

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("docinsight")

app = FastAPI(
    title="DocInsight API",
    description="Upload and query documents (.pdf, .docx, .eml) using GPT-4o",
    version="1.0"
)

# --- FRONTEND SERVING (robust) ---
# Try several likely locations for the frontend folder (case-insensitive)
_possible_dirs = [
    os.path.join(os.path.dirname(__file__), "frontend"),
    os.path.join(os.path.dirname(__file__), "Frontend"),
    os.path.join(os.path.dirname(__file__), "..", "frontend"),
    os.path.join(os.path.dirname(__file__), "..", "Frontend"),
    os.path.join(os.getcwd(), "doc-insight-api", "frontend"),
    os.path.join(os.getcwd(), "doc-insight-api", "Frontend"),
    os.path.join(os.getcwd(), "frontend"),
    os.path.join(os.getcwd(), "Frontend"),
]

frontend_path: Optional[str] = None
for p in _possible_dirs:
    if os.path.isdir(p):
        frontend_path = os.path.abspath(p)
        break

if frontend_path:
    logger.info(f"Serving frontend from: {frontend_path}")
    # mount static files under /static to avoid clobbering other routes
    app.mount("/static", StaticFiles(directory=frontend_path), name="static")

    # serve index at root
    @app.get("/", include_in_schema=False)
    async def serve_index():
        index_file = os.path.join(frontend_path, "index.html")
        if os.path.isfile(index_file):
            return FileResponse(index_file, media_type="text/html")
        raise HTTPException(status_code=404, detail="Index file not found")
else:
    logger.warning("Frontend folder not found; static files will not be served. Expected a 'frontend' folder.")

# Allow CORS (for local dev and separate static host). Tighten in production.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OpenAI client (new-style client)
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_KEY:
    logger.warning("OPENAI_API_KEY not found in environment. OpenAI calls will fail until it's set.")
client = OpenAI(api_key=OPENAI_KEY) if OPENAI_KEY else None

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def extract_pdf_text(path: str) -> str:
    with fitz.open(path) as doc:
        pages = []
        for page in doc:
            pages.append(page.get_text())
        return "\n".join(pages)

def extract_docx_text(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_eml_text(path: str) -> str:
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email.policy.default)
    parts = []
    if msg.is_multipart():
        for part in msg.walk():
            # prefer plain text
            if part.get_content_type() == "text/plain":
                try:
                    parts.append(part.get_content())
                except Exception:
                    parts.append(part.get_payload(decode=True).decode(errors="ignore"))
    else:
        try:
            parts.append(msg.get_content())
        except Exception:
            parts.append(msg.get_payload(decode=True).decode(errors="ignore"))
    return "\n".join(parts)

def count_tokens(text: str, model: str = "gpt-4o") -> int:
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except Exception:
        # fallback heuristic
        return max(1, len(text) // 4)

# In-memory conversation store for quick chat testing.
# { doc_id (str) : [ {"role":"user"|"assistant", "content": "..."} , ... ] }
CONVERSATIONS: Dict[str, List[Dict[str, str]]] = {}

class ChatRequest(BaseModel):
    doc_id: str
    message: str

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """
    Uploads a file (pdf/docx/eml), extracts text, stores in DB, returns doc_id
    """
    try:
        suffix = os.path.splitext(file.filename)[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        logger.info(f"Saved uploaded file to: {tmp_path} (suffix={suffix})")

        if suffix == ".pdf":
            text = extract_pdf_text(tmp_path)
        elif suffix == ".docx":
            text = extract_docx_text(tmp_path)
        elif suffix == ".eml":
            text = extract_eml_text(tmp_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        doc_id = str(uuid.uuid4())

        # Save minimal document row to DB
        db_doc = Document(id=doc_id, filename=file.filename, filetype=suffix, content=text)
        db.add(db_doc)
        db.commit()

        # initialize conversation store
        CONVERSATIONS[doc_id] = []

        logger.info(f"Document stored with id: {doc_id} (length={len(text)} chars)")
        return {"message": "File uploaded and parsed", "doc_id": doc_id}

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("File upload failed")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat")
async def chat(request: ChatRequest, db: Session = Depends(get_db)):
    """
    Chat endpoint:
    Request JSON: { "doc_id": "<uuid>", "message":"your question" }
    Response JSON: { "answer": "...", "history": [ {role, content}, ... ] }
    """
    if client is None:
        raise HTTPException(status_code=500, detail="OpenAI client not configured (OPENAI_API_KEY missing)")

    try:
        doc_id = request.doc_id
        message = request.message

        db_doc = db.query(Document).filter(Document.id == str(doc_id)).first()
        if not db_doc:
            raise HTTPException(status_code=404, detail="Document not found")

        full_text = db_doc.content or ""

        # Truncate document to max token length (safe default)
        max_tokens_for_document = 6000
        try:
            encoding = tiktoken.encoding_for_model("gpt-4o")
            tokens = encoding.encode(full_text)
            truncated_text = encoding.decode(tokens[:max_tokens_for_document])
        except Exception:
            truncated_text = full_text[:100000]  # fallback

        # Build conversation history text
        history_items = CONVERSATIONS.get(doc_id, [])
        history_text = "\n".join(
            [f"User: {m['content']}" if m['role'] == 'user' else f"Assistant: {m['content']}" for m in history_items]
        )

        # Prompt instructing the model to answer ONLY from the document and in one-sentence JSON.
        prompt = (
            "You are a helpful assistant. Use ONLY the provided document content and conversation history to answer the user's question.\n\n"
            "Answer in exactly one clear, concise sentence and return only a JSON object: {\"answer\": \"...\"}\n\n"
            "--- Document Text ---\n"
            f"{truncated_text}\n"
            "----------------------\n\n"
            "--- Conversation so far ---\n"
            f"{history_text}\n"
            "---------------------------\n\n"
            f"User Question: {message}\n\n"
            "Respond in this JSON format only:\n"
            '{"answer": "your single sentence answer here"}'
        )

        logger.info("Sending prompt to OpenAI (truncated document length=%d chars)", len(truncated_text))

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You answer strictly in valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=200
        )

        # parse model response
        content = response.choices[0].message.content.strip()
        logger.info(f"OpenAI raw response: {content[:1000]}")

        # remove code fences if present
        if content.startswith("```json"):
            content = content.replace("```json", "").replace("```", "").strip()
        elif content.startswith("```"):
            content = content.replace("```", "").strip()

        try:
            parsed = json.loads(content)
        except json.JSONDecodeError:
            logger.error("Invalid JSON returned by model. Returning raw content wrapped in answer.")
            parsed = {"answer": content}

        answer = parsed.get("answer", parsed.get("Answer", parsed.get("response", "")))

        # Update in-memory conversation
        CONVERSATIONS.setdefault(doc_id, []).append({"role": "user", "content": message})
        CONVERSATIONS.setdefault(doc_id, []).append({"role": "assistant", "content": answer})

        # Log query in DB if you have QueryLog
        try:
            db.add(QueryLog(query_text=message, response_json=json.dumps(parsed)))
            db.commit()
        except Exception:
            db.rollback()

        return {"answer": answer, "history": CONVERSATIONS.get(doc_id, [])}

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("OpenAI query failed")
        raise HTTPException(status_code=500, detail="OpenAI query failed")

# Document listing + delete endpoints (unchanged logic)
@app.get("/documents/")
def list_documents(db: Session = Depends(get_db)):
    try:
        docs = db.query(Document).all()
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "filetype": doc.filetype,
                "upload_time": getattr(doc, "uploaded_at", None).isoformat() if getattr(doc, "uploaded_at", None) else None,
                "preview": doc.content[:100] + "..." if doc.content else ""
            }
            for doc in docs
        ]
    except Exception:
        logger.exception("Failed to list documents")
        raise HTTPException(status_code=500, detail="Unable to fetch documents")

@app.delete("/delete/{doc_id}")
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        db.delete(doc)
        db.commit()
        if doc_id in CONVERSATIONS:
            del CONVERSATIONS[doc_id]
        return {"message": f"Document {doc_id} deleted"}
    except Exception:
        logger.exception("Failed to delete document")
        raise HTTPException(status_code=500, detail="Unable to delete document")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)